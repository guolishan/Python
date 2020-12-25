import sys
import docx
import os

'''
遍历目录在wrod里插入图片。
'''

def file_name_walk(file_dir, kwword):
    """
    返回要求文件的绝对路径(短边右边测试报告)
    :param file_dir: 目录
    :param kwword: 文件编号
    :return:
    """
    for root, dirs, files in os.walk(file_dir):
        # print("root", root)  # 当前目录路径
        # print("dirs", dirs)  # 当前路径下所有子目录
        # print("files", files)  # 当前路径下所有非目录子文件
        for f in files:
            if kwword in f:
                return root + "\\" + f


# 拿到图片的名字，并按逗号分隔
def photo_name_walk(photo_dir,file_dir, suffix='.png'):
    """
    从图片名字里拿到能找到word的编号（或者名字）
    :param photo_dir: 图片路径
    :param file_dir: 要插入wrod文件的路径
    :param suffix: 文件后缀
    :return:
    """
    for root, dirs, files in os.walk(photo_dir):
        for f in files:
            if os.path.splitext(f)[1] == suffix:
                photo_name = os.path.splitext(f)[0]
                # 对photo按照逗号进行分割
                file_number = photo_name.split(',')[0]  # 利用file_number找到对应的word文档
                position_number = photo_name.split(',')[1]  # 利用position找到文档要插入的地方
                # photo 的绝对路径
                photo_dit_path = root + "\\" + f
                # 找到word的绝对路径
                file_dir_path = file_name_walk(file_dir, file_number)
                # 在word里插入图片
                word_main(file_dir_path, photo_dit_path, position_number)


def word_main(file_dir_path, photo_dit_path, position_number):
    '''
    在word里插入图片
    :param file_dir_path: 文件绝对路径
    :param photo_dit_path: 图片绝对路径
    :param position_number: 要定位的地方，在该word中找position_number这串字符串
    :return:
    '''
    doc = docx.Document(file_dir_path)
    for i, p in enumerate(doc.paragraphs):  # 遍历所有的段落
        print(str(i) + ":"+ str(p.text))
        if len(p.text) != 0:

            for i in range(len(p.runs)):  # p.runs代表p这个段落下所有文字的列表
                print(str(i)+':::::')
                print(p.runs[i].text)  # 当打印时，发现p.runs把段落自动分解了
        if position_number in p.text:
            p.runs[-1].add_break()  # 添加一个折行
            p.runs[-1].add_picture(photo_dit_path)  # 在runs的最后一段文字后添加图片
            # os.remove(photo_dit_path)
            doc.save(file_dir_path)  # 保存文件
            break


file_dir = sys.argv[1]
# print('file_dir', file_dir)
photo_dir = sys.argv[2]
# print('photo_dir', photo_dir)
photo_name_walk(photo_dir, file_dir)