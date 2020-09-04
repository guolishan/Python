import os
import sys

import docx
from PyQt5.QtGui import QTextCursor, QIcon
from PyQt5.QtWidgets import QWidget, QApplication, QGroupBox, QPushButton, QLabel, QHBoxLayout, QVBoxLayout, \
    QGridLayout, QLineEdit, QTextEdit, QFileDialog, QMessageBox


class ImgCut(QWidget):

    def __init__(self):
        super(ImgCut, self).__init__()

        # 存储位置
        self.forder_label = QLabel('文档位置：')
        self.forder_text = QLineEdit()
        self.forder_text.setDisabled(True)
        self.forder_btn = QPushButton('选择文档位置…')

        # 要替换的关键词
        self.foot_label = QLabel('关键词：')
        self.foot_text = QLineEdit()
        self.foot_text.setPlaceholderText('多个关键词之间使用英文半角,分割')

        # 替换后的关键词
        self.nums_label = QLabel('替换后：')
        self.nums_text = QLineEdit()
        self.nums_text.setPlaceholderText('与关键词一致，一一对应')




        # 确认按钮
        self.submit_btn = QPushButton('开始替换')
        self.submit_btn.setStyleSheet("QPushButton{padding:20px 4px}")

        # 图片label
        self.img_label = QLabel()

        # 复制按钮
        self.copy_btn = QPushButton('复制代码')
        self.copy_btn.setStyleSheet("QPushButton{padding:20px 4px}")
        # 显示生成的文件html
        self.res_teatarea = QTextEdit()

        self.initUi()

    def initUi(self):
        self.createGridGroupBox()
        # self.creatVboxGroupBox()
        self.creatFormGroupBox()
        mainLayout = QVBoxLayout()
        hboxLayout = QHBoxLayout()
        # hboxLayout.addStretch()
        hboxLayout.addWidget(self.gridGroupBox)
        # hboxLayout.addWidget(self.vboxGroupBox)
        mainLayout.addLayout(hboxLayout)
        mainLayout.addWidget(self.formGroupBox)
        self.setLayout(mainLayout)
        # 禁止最大化
        # self.setFixedSize(self.width(), self.height())

    # 参数区域
    def createGridGroupBox(self):
        self.gridGroupBox = QGroupBox("基本参数")
        layout = QGridLayout()

        # 点击选择保存路径按钮
        self.forder_btn.clicked.connect(self.savePath)
        # 点击提交按钮
        self.submit_btn.clicked.connect(self.submit)
        layout.setSpacing(10)
        # 网格布局
        layout.addWidget(self.forder_label, 1, 0)
        layout.addWidget(self.forder_text, 1, 1)
        layout.addWidget(self.forder_btn, 1, 2)

        layout.addWidget(self.foot_label, 2, 0)
        layout.addWidget(self.foot_text, 2, 1, 1, 2)
        layout.addWidget(self.nums_label, 3, 0)
        layout.addWidget(self.nums_text, 3, 1, 1, 2)

        layout.addWidget(self.submit_btn, 4, 0, 1, 3)
        layout.setColumnStretch(1, 10)
        self.gridGroupBox.setLayout(layout)
        self.setWindowTitle('word替换')
        self.setWindowIcon(QIcon(r'E:\Python_Project\Python\00_项目文件库\1_UI\logo.ico'))


    # 代码区域
    def creatFormGroupBox(self):
        self.formGroupBox = QGroupBox("替换结果")
        layout = QGridLayout()

        layout.addWidget(self.res_teatarea, 1, 0)
        # layout.addWidget(self.copy_btn, 2, 0)
        # 点击选择保存路径按钮
        self.copy_btn.clicked.connect(self.copyText)

        self.formGroupBox.setLayout(layout)

    # 显示消息
    def showMsg(self, tit, content, icon=3):
        box = QMessageBox(QMessageBox.Question, tit, content)
        # 设置左上角消息框图标
        box.setWindowIcon(QIcon(r'E:\Python_Project\Python\00_项目文件库\1_UI\logo.ico'))

        # 添加按钮，可用中文
        yes = box.addButton('确定', QMessageBox.YesRole)
        # 设置消息框中内容前面的图标
        box.setIcon(icon)
        # 显示该问答框
        box.exec()

    # 选择保存文件夹
    def savePath(self):
        path = QFileDialog.getExistingDirectory(self, "请选择您要保存的位置")
        # 判断选择的文件是否存在
        if os.path.exists(path):
            # 将保存url放入路径文本框中
            self.forder_text.setText(path)
        else:
            self.showMsg('错误', '您选择的保存路径不存在，请重新选择！')

    # 提交替换
    def submit(self):
        forder = self.forder_text.text()
        if not forder:
            self.showMsg('错误', '您还没有选择文档所在目录')
            return False

        # 获取要替换的关键词
        before_words = self.foot_text.text()
        # 判断是否填写要替换的关键词
        if before_words:
            before_words = before_words.split(',')
        else:
            self.showMsg('错误', '请填写要替换的关键词')

        # 获取替换后的词
        after_words = self.nums_text.text().split(',')

        # 获取文件列表
        files = self.get_files(forder)

        for file in files:

            doc = docx.Document(file)

            # 判断替换前后关键词长度是否一致(一致就按位置替换关键词，否则使用第一个被替换的额关键词进行替换)
            if len(before_words) == len(after_words):
                for index, word in enumerate(before_words):
                    self.replace_word(doc, word, after_words[index])
            else:
                for word in before_words:
                    self.replace_word(doc, word, after_words[0])
            # 保存word
            doc.save("{}/{}".format(forder, file.split("/")[-1]))
            # 代码框光标移动至末尾
            self.res_teatarea.moveCursor(QTextCursor.End)
            # 每次打印的替换结果
            shtml = "{}替换完成\n".format(file)
            # 将代码粘贴至代码框光标位置
            self.res_teatarea.insertPlainText(shtml)


	# 获取docx文件列表
    def get_files(self, forder):
        files = []
        for file in os.listdir(forder):
            if file.endswith(".docx"): #排除文件夹内的其它干扰文件，只获取word文件
                files.append(forder+'/'+file)
        return files

    # 替换文档关键词
    def replace_word(self, doc, before_words, after_words):
        '''此函数用于批量替换合同中需要替换的信息
        doc:文件
        old_info和new_info：原文字和需要替换的新文字
        '''
        # 读取段落中的所有run，找到需替换的信息进行替换
        for para in doc.paragraphs:  #
            for run in para.runs:
                run.text = run.text.replace(before_words, after_words)  # 替换信息
        # 读取表格中的所有单元格，找到需替换的信息进行替换
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell.text = cell.text.replace(before_words, after_words)  # 替换信息


    # 复制代码，已废弃
    def copyText(self):
        # 获取代码框内容
        data = self.res_teatarea.toPlainText()
        # 判断代码框是否有内容
        if data:
            # 如果有内容将内容添加至剪贴板
            clipboard = QApplication.clipboard()
            clipboard.setText(data)
            self.showMsg('信息', '内容以成功复制到剪贴板', 1)
        else:
            self.showMsg('错误', '代码为空，没有代码可以复制')



if __name__ == '__main__':
    app = QApplication(sys.argv)
    ex = ImgCut()
    ex.show()
    sys.exit(app.exec_())